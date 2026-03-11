import streamlit as st
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import io
from PIL import Image
from datetime import date

# --- AYARLAR ---
SAVE_DIR = "TEKLİFLER"
TEMP_DIR = "GECICI_RESIMLER"
LOGO_PATH = "logo.png"

for klasor in [SAVE_DIR, TEMP_DIR]:
    if not os.path.exists(klasor): os.makedirs(klasor)

def dosya_adi_yap(metin):
    replacements = {'ı':'i','İ':'I','ş':'s','Ş':'S','ğ':'g','Ğ':'G','ü':'u','Ü':'U','ö':'o','Ö':'O','ç':'c','Ç':'C', 'â':'a', 'Â':'A'}
    metin = str(metin)
    for k, v in replacements.items():
        metin = metin.replace(k, v)
    return metin.replace(" ", "_")

# --- ARAYÜZ ---
st.set_page_config(page_title="Alpinetech Teklif Hazırlama", layout="wide")
st.title("🏔️ Alpinetech")

# --- HAFIZA VE SABİT VERİTABANI ---
if 'kategoriler' not in st.session_state:
    st.session_state.kategoriler = []
    
if 'sartlar' not in st.session_state:
    st.session_state.sartlar = [
        "Ulaşım ve konaklama giderleri teklif bedeline dâhildir.",
        "Fiyatlarımıza KDV (%20) dâhil değildir.",
        "Ödeme Koşulları peşin.",
        "İş bu teklif ve şartlar, verildiği tarihten itibaren 15 (on beş) takvim günü boyunca geçerlidir.",
        "Montaj için gerekli saha düzenlemeleri ve makine temin ve çalışmaları müşteriye aittir.",
        "Gerekli izinlerin alınması müşteriye aittir.",
    ]

if 'musteriler' not in st.session_state:
    st.session_state.musteriler = {
        "Murat Bey'in Dikkatine": {"tel": "+90 533 741 38 60", "mail": "alp.orme1@gmail.com"},
        "Örnek Firma A.Ş.": {"tel": "+90 555 111 22 33", "mail": "info@ornek.com"}
    }

# --- 1. ÜST BİLGİLER VE MÜŞTERİ/VEREN BİLGİLERİ ---
with st.expander("📄 Teklif, İletişim ve KDV Ayarları", expanded=True):
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("**📌 Teklif Detayları**")
        t_baslik1 = st.text_input("Başlık (1. Satır)", "40 Metre Yürüyen Bant Fiyat Teklifi")
        t_baslik2 = st.text_input("Başlık (2. Satır)", "")
        t_tarih = st.text_input("Tarih", date.today().strftime("%d.%m.%Y"))
        para_birimi = st.selectbox("Para Birimi", ["EUR", "TL", "USD"])
        
        st.markdown("**💰 Vergi Ayarları**")
        k1, k2 = st.columns(2)
        with k1:
            kdv_ekle = st.checkbox("KDV Hesapla", value=True)
        with k2:
            kdv_oran = st.number_input("KDV Oranı (%)", min_value=0, value=20, step=1)

    with c2:
        st.markdown("**🏢 Müşteri Bilgileri**")
        secenekler = ["-- Yeni Müşteri Gir --"] + list(st.session_state.musteriler.keys())
        secilen_musteri = st.selectbox("📇 Kayıtlı Müşteri Seç", secenekler)
        
        if secilen_musteri != "-- Yeni Müşteri Gir --":
            def_ad = secilen_musteri
            def_tel = st.session_state.musteriler[secilen_musteri].get("tel", "")
            def_mail = st.session_state.musteriler[secilen_musteri].get("mail", "")
        else:
            def_ad, def_tel, def_mail = "", "", ""
            
        m_ad = st.text_input("Müşteri Firma/Kişi", value=def_ad)
        m_tel = st.text_input("Müşteri Tel", value=def_tel)
        m_mail = st.text_input("Müşteri E-Posta", value=def_mail)
        
        if st.button("💾 Müşteriyi Hafızaya Al", use_container_width=True):
            if m_ad:
                st.session_state.musteriler[m_ad] = {"tel": m_tel, "mail": m_mail}
                st.success(f"{m_ad} oturum hafızasına eklendi!")
                st.rerun()

    with c3:
        st.markdown("**🧑‍💻 Teklifi Veren Bilgileri**")
        v_ad = st.text_input("İsim Soyisim", "Kürşad Nuri Örme")
        v_tel = st.text_input("Telefon", "+90 541 575 21 70")
        v_mail = st.text_input("E-Posta", "info@alpinetech.com.tr")

st.divider()

# --- 2. KATEGORİ, ADET VE MADDE EKLEME ---
st.subheader("📋 İş Kalemleri")
with st.form("kategori_ekle", clear_on_submit=True):
    col_k1, col_k2 = st.columns([3, 1])
    with col_k1:
        kat_adi = st.text_input("Kategori Adı (Örn: 40 Metre Yürüyen Bant Sistemi)")
        kat_aciklama = st.text_area("Maddeler (Alt alta yazabilirsiniz)")
        yuklenenler = st.file_uploader("Ürün Görselleri (Yan yana dizilecek)", type=['jpg','png','jpeg'], accept_multiple_files=True)
    with col_k2:
        kat_adet = st.number_input("Adet", min_value=1.0, value=1.0, step=1.0)
        kat_fiyat = st.number_input("Birim Fiyat", min_value=0.0)
        st.markdown("<br>", unsafe_allow_html=True)
        ekle_btn = st.form_submit_button("➕ Kategori Ekle")
        
    if ekle_btn and kat_adi:
        adet_val = int(kat_adet) if kat_adet.is_integer() else kat_adet
        resim_yollari = []
        if yuklenenler:
            for i, dosya in enumerate(yuklenenler):
                try:
                    yol = os.path.join(TEMP_DIR, f"img_{len(st.session_state.kategoriler)}_{i}.jpg")
                    img = Image.open(dosya).convert("RGB")
                    img.save(yol, "JPEG")
                    resim_yollari.append(yol)
                except: pass
                    
        st.session_state.kategoriler.append({
            "baslik": kat_adi, 
            "aciklama": kat_aciklama, 
            "adet": adet_val,
            "fiyat": kat_fiyat,
            "toplam": adet_val * kat_fiyat,
            "resimler": resim_yollari
        })

if st.session_state.kategoriler:
    for i, k in enumerate(st.session_state.kategoriler):
        c_disp1, c_disp2 = st.columns([9, 1])
        with c_disp1:
            st.markdown(f"**<span style='color:#2E74B5'>{k['baslik']}</span>**", unsafe_allow_html=True)
            if k['resimler']:
                cols = st.columns(min(len(k['resimler']), 4))
                for idx, r in enumerate(k['resimler']):
                    cols[idx % 4].image(r, width=150)
            st.text(k['aciklama'])
            st.write(f"*Adet: {k['adet']} | Birim Fiyat: {k['fiyat']:,.0f} {para_birimi} | **Toplam: {k['toplam']:,.0f} {para_birimi}***")
        with c_disp2:
            if st.button("🗑️ Sil", key=f"kat_sil_{i}"):
                st.session_state.kategoriler.pop(i); st.rerun()
        st.write("---")

# --- 3. GENEL ŞARTLAR ---
st.subheader("⚖️ Genel Şartlar")
with st.form("sart_ekle", clear_on_submit=True):
    c_sart1, c_sart2 = st.columns([4, 1])
    yeni_sart = c_sart1.text_input("Yeni Şart Ekle")
    if c_sart2.form_submit_button("➕ Şart Ekle") and yeni_sart:
        st.session_state.sartlar.append(yeni_sart)

if st.session_state.sartlar:
    for i, sart in enumerate(st.session_state.sartlar):
        sc1, sc2 = st.columns([9, 1])
        sc1.write(f"- {sart}")
        if sc2.button("🗑️ Sil", key=f"sart_sil_{i}"):
            st.session_state.sartlar.pop(i); st.rerun()

st.divider()

# --- 4. DOSYA KAYIT VE FORMAT SEÇİMİ ---
st.subheader("💾 Kayıt Ayarları")
col_f1, col_f2 = st.columns([2, 1])
with col_f1: dosya_adi = st.text_input("Kaydedilecek Dosya Adı", value="Yeni_Teklif_Dosyasi")
with col_f2: format_secimi = st.radio("Dosya Formatı", ["PDF", "Word (.docx)"], horizontal=True)

if st.button("🚀 TEKLİFİ OLUŞTUR VE İNDİR", type="primary"):
    if not st.session_state.kategoriler:
        st.error("Lütfen en az bir kategori ekleyin!")
    elif not dosya_adi:
        st.error("Lütfen bir dosya adı girin!")
    else:
        ara_toplam = sum(k['toplam'] for k in st.session_state.kategoriler)
        kdv_tutari = ara_toplam * (kdv_oran / 100) if kdv_ekle else 0
        genel_toplam = ara_toplam + kdv_tutari

        # --- PDF OLUŞTURMA ---
        if format_secimi == "PDF":
            pdf = FPDF()
            pdf.add_page()
            
            font_adi = 'Arial'
            if os.path.exists('arial.ttf') and os.path.exists('arialbd.ttf'):
                pdf.add_font('ArialTR', '', 'arial.ttf', uni=True)
                pdf.add_font('ArialTR', 'B', 'arialbd.ttf', uni=True)
                font_adi = 'ArialTR'
            elif os.path.exists(r'C:\Windows\Fonts\arial.ttf'):
                pdf.add_font('ArialTR', '', r'C:\Windows\Fonts\arial.ttf', uni=True)
                pdf.add_font('ArialTR', 'B', r'C:\Windows\Fonts\arialbd.ttf', uni=True)
                font_adi = 'ArialTR'
                
            def p(metin):
                if font_adi == 'ArialTR': return str(metin)
                rep = {'ı':'i','İ':'I','ş':'s','Ş':'S','ğ':'g','Ğ':'G','ü':'u','Ü':'U','ö':'o','Ö':'O','ç':'c','Ç':'C', 'â':'a', 'Â':'A', '•':'-', '€':'EUR', '₺':'TL'}
                m = str(metin)
                for k, v in rep.items(): m = m.replace(k, v)
                return m.encode('latin-1', 'ignore').decode('latin-1')
            
            BLUE, GREEN, BLACK = (46, 116, 181), (112, 173, 71), (0, 0, 0)
            
            if os.path.exists(LOGO_PATH): pdf.image(LOGO_PATH, x=130, y=14, w=72)
                
            pdf.set_font(font_adi, 'B', 9); pdf.set_xy(10, 15)
            pdf.cell(0, 4.5, p("Alpinetech Mühendislik Makine"), ln=1)
            pdf.cell(0, 4.5, p("Sanayi ve Ticaret Limited Şirketi"), ln=1)
            pdf.cell(0, 4.5, p("30 Ağustos Zafer Mahallesi 520."), ln=1)
            pdf.cell(0, 4.5, p("Cadde 4C/A Nilüfer Bursa"), ln=1)
            pdf.set_text_color(*GREEN); pdf.cell(0, 4.5, "www.alpinetech.com.tr", ln=1)
            pdf.set_text_color(*BLACK); pdf.cell(0, 4.5, "info@alpinetech.com.tr", ln=1)

            pdf.ln(12); pdf.set_font(font_adi, 'B', 16); pdf.cell(0, 8, p(f"TEKLİF: {t_baslik1}"), ln=1)
            pdf.set_font(font_adi, 'B', 14); pdf.cell(100, 8, p(t_baslik2), ln=0)
            pdf.set_font(font_adi, '', 11); pdf.cell(90, 8, p(t_tarih), ln=1, align='R')
            pdf.set_line_width(0.3); pdf.line(10, pdf.get_y()+2, 200, pdf.get_y()+2); pdf.ln(5)

            pdf.set_font(font_adi, 'B', 9); y_info = pdf.get_y()
            pdf.set_xy(10, y_info); pdf.cell(15, 5, p("Müşteri"), ln=0); pdf.set_font(font_adi, ''); pdf.cell(70, 5, p(m_ad), ln=1)
            pdf.set_font(font_adi, 'B'); pdf.cell(15, 5, "Telefon", ln=0); pdf.set_font(font_adi, ''); pdf.cell(70, 5, p(m_tel), ln=1)
            pdf.set_font(font_adi, 'B'); pdf.cell(15, 5, "E-Posta", ln=0); pdf.set_font(font_adi, ''); pdf.cell(70, 5, p(m_mail), ln=1)
            
            pdf.set_font(font_adi, 'B', 9); pdf.set_xy(115, y_info)
            pdf.cell(25, 5, "Teklifi veren", ln=0); pdf.set_font(font_adi, ''); pdf.cell(50, 5, p(v_ad), ln=1)
            pdf.set_xy(115, y_info+5); pdf.set_font(font_adi, 'B', 9); pdf.cell(25, 5, "Telefon", ln=0); pdf.set_font(font_adi, ''); pdf.cell(50, 5, p(v_tel), ln=1)
            pdf.set_xy(115, y_info+10); pdf.set_font(font_adi, 'B', 9); pdf.cell(25, 5, "E-posta", ln=0); pdf.set_text_color(*GREEN); pdf.cell(50, 5, p(v_mail), ln=1)
            pdf.set_text_color(*BLACK); pdf.ln(3); pdf.line(10, pdf.get_y(), 200, pdf.get_y()); pdf.ln(5)

            for k in st.session_state.kategoriler:
                if pdf.get_y() > 240: pdf.add_page() # Kalem sığmıyorsa atla
                
                pdf.set_text_color(*BLUE); pdf.set_font(font_adi, '', 14); pdf.cell(0, 8, p(k['baslik']), ln=1)
                pdf.set_draw_color(*BLUE); pdf.set_line_width(0.2); pdf.line(10, pdf.get_y(), 200, pdf.get_y()); pdf.set_draw_color(*BLACK) 
                
                # Açıklama
                pdf.set_text_color(*BLACK); pdf.set_font(font_adi, 'B', 9); pdf.ln(1)
                pdf.multi_cell(0, 5, p(k['aciklama']))
                
                # FİYAT KUTUSU (RESİMDEN ÖNCEYE ALINDI)
                pdf.ln(2); pdf.set_font(font_adi, '', 9)
                kutu_fiyat = f"{k['toplam']:,.0f}".replace(",", ".")
                pdf.cell(0, 6, p(f"Fiyat : {kutu_fiyat} {para_birimi}"), border=1, align='R', ln=1)
                
                # RESİMLER (FİYATTAN SONRA, YAN YANA)
                if k['resimler']:
                    pdf.ln(4)
                    if pdf.get_y() > 210: pdf.add_page() # Sadece resim sığmıyorsa resimleri diğer sayfaya at
                    
                    y_start = pdf.get_y()
                    img_height = 0
                    for idx, r in enumerate(k['resimler'][:2]): # Max 2 resim yan yana
                        if os.path.exists(r):
                            pdf.image(r, x=10 + (idx * 95), y=y_start, w=85)
                            img_height = 60 # Ortalama resim yüksekliği
                    if img_height > 0:
                        pdf.set_y(y_start + img_height + 5)
                else:
                    pdf.ln(2)

            # --- GENEL TOPLAM TABLOSU VE ŞARTLAR (DİNAMİK YERLEŞİM) ---
            # Tablo ve şartların sayfaya sığıp sığmadığını kontrol ediyoruz
            gereken_bosluk = 20 + (len(st.session_state.kategoriler) * 5) + (20 if kdv_ekle else 10)
            if pdf.get_y() + gereken_bosluk > 275:
                pdf.add_page()
            else:
                pdf.ln(10)

            pdf.set_text_color(*BLUE); pdf.set_font(font_adi, '', 12)
            pdf.cell(40, 6, "", ln=0); pdf.cell(50, 6, "Genel Toplam", ln=0, align='R'); pdf.cell(20, 6, "Adet", ln=0, align='C'); pdf.cell(35, 6, "Fiyat", ln=0, align='C'); pdf.cell(45, 6, "Toplam", ln=1, align='C')
            pdf.set_draw_color(*BLUE); pdf.line(50, pdf.get_y(), 200, pdf.get_y()); pdf.set_draw_color(*BLACK); pdf.set_text_color(*BLACK)
            
            pdf.set_font(font_adi, 'B', 9); pdf.ln(2)
            for k in st.session_state.kategoriler:
                f_str = f"{k['fiyat']:,.0f}".replace(",", "."); t_str = f"{k['toplam']:,.0f}".replace(",", ".")
                pdf.cell(40, 5, "", ln=0); pdf.cell(50, 5, p(k['baslik']), ln=0, align='R'); pdf.cell(20, 5, str(k['adet']), ln=0, align='C')
                pdf.cell(35, 5, p(f"{f_str} {para_birimi}"), ln=0, align='C'); pdf.cell(45, 5, p(f"{t_str} {para_birimi}"), ln=1, align='C')

            pdf.ln(2)
            if kdv_ekle:
                pdf.cell(110, 5, "", ln=0); pdf.cell(35, 5, "Toplam", ln=0, align='R'); pdf.cell(45, 5, p(f"{ara_toplam:,.0f} {para_birimi}"), ln=1, align='C')
                pdf.cell(110, 5, "", ln=0); pdf.cell(35, 5, p(f"KDV(%{kdv_oran})"), ln=0, align='R'); pdf.cell(45, 5, p(f"{kdv_tutari:,.0f} {para_birimi}"), ln=1, align='C')
            pdf.cell(110, 5, "", ln=0); pdf.cell(35, 5, "Genel Toplam", ln=0, align='R'); pdf.cell(45, 5, p(f"{genel_toplam:,.0f} {para_birimi}"), ln=1, align='C')

            # Şartlar bölümünün sığıp sığmadığını kontrol et
            if st.session_state.sartlar:
                gereken_sart_boslugu = 15 + (len(st.session_state.sartlar) * 6)
                if pdf.get_y() + gereken_sart_boslugu > 275:
                    pdf.add_page()
                else:
                    pdf.ln(10)
                    
                pdf.set_text_color(*BLUE); pdf.set_font(font_adi, '', 14); pdf.cell(0, 8, p("Genel Şartlar"), ln=1)
                pdf.set_draw_color(*BLUE); pdf.line(10, pdf.get_y(), 200, pdf.get_y()); pdf.set_draw_color(*BLACK)
                pdf.ln(2); pdf.set_text_color(*BLACK); pdf.set_font(font_adi, '', 9)
                for sart in st.session_state.sartlar:
                    pdf.cell(5, 5, "-", ln=0); pdf.multi_cell(0, 5, p(sart))

            pdf.set_auto_page_break(False); pdf.set_y(-15); pdf.set_font(font_adi, '', 8); pdf.cell(0, 10, "1", 0, 0, 'R')
            
            safe_name = dosya_adi_yap(dosya_adi) + ".pdf"
            out_path = os.path.join(SAVE_DIR, safe_name)
            pdf.output(out_path)
            st.success(f"✅ PDF '{safe_name}' olarak kaydedildi!")
            with open(out_path, "rb") as f: st.download_button("📩 PDF'i İndir", f, file_name=safe_name)

        # --- WORD OLUŞTURMA (.docx) ---
        else:
            doc = Document()
            title_p = doc.add_paragraph()
            t1 = title_p.add_run(f"TEKLİF: {t_baslik1}"); t1.bold = True; t1.font.size = Pt(16)
            doc.add_paragraph(f"{t_baslik2}").bold = True
            doc.add_paragraph(f"Tarih: {t_tarih}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_heading('Firma ve İletişim Bilgileri', level=2)
            doc.add_paragraph(f"Müşteri Firma/Kişi: {m_ad}\nTel: {m_tel}\nE-Posta: {m_mail}")
            doc.add_paragraph(f"Teklifi Veren: {v_ad}\nTel: {v_tel}\nE-Posta: {v_mail}")
            
            doc.add_heading('İş Kalemleri', level=2)
            for k in st.session_state.kategoriler:
                p = doc.add_paragraph()
                r = p.add_run(k['baslik'])
                r.bold = True; r.font.color.rgb = RGBColor(46, 116, 181); r.font.size = Pt(12)
                
                # Açıklama
                doc.add_paragraph(k['aciklama'])
                
                # Fiyat Kutusu (Resimden önce)
                p_fiyat = doc.add_paragraph(f"Fiyat: {k['toplam']:,.0f} {para_birimi}")
                p_fiyat.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Resimler
                if k['resimler']:
                    table = doc.add_table(rows=1, cols=min(len(k['resimler']), 2))
                    for idx, img_path in enumerate(k['resimler'][:2]):
                        if os.path.exists(img_path):
                            cell = table.cell(0, idx)
                            cell.paragraphs[0].add_run().add_picture(img_path, width=Inches(3.0))
            
            doc.add_paragraph("\n")
            doc.add_heading('Genel Toplam Özeti', level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'İş Kalemi'; hdr_cells[1].text = 'Adet'; hdr_cells[2].text = 'Birim Fiyat'; hdr_cells[3].text = 'Toplam'
            
            for k in st.session_state.kategoriler:
                row_cells = table.add_row().cells
                row_cells[0].text = k['baslik']; row_cells[1].text = str(k['adet']); row_cells[2].text = f"{k['fiyat']:,.0f} {para_birimi}"; row_cells[3].text = f"{k['toplam']:,.0f} {para_birimi}"
            
            doc.add_paragraph("\n")
            if kdv_ekle:
                doc.add_paragraph(f"Ara Toplam: {ara_toplam:,.0f} {para_birimi}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph(f"KDV (%{kdv_oran}): {kdv_tutari:,.0f} {para_birimi}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            p_total = doc.add_paragraph(f"GENEL TOPLAM: {genel_toplam:,.0f} {para_birimi}")
            p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p_total.runs[0].bold = True

            if st.session_state.sartlar:
                doc.add_heading('Genel Şartlar', level=2)
                for sart in st.session_state.sartlar: doc.add_paragraph(sart, style='List Bullet')

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            safe_name = dosya_adi_yap(dosya_adi) + ".docx"
            st.success(f"✅ Word dosyası '{safe_name}' olarak hazırlandı!")
            st.download_button("📩 Word Dosyasını İndir", buffer, file_name=safe_name)

